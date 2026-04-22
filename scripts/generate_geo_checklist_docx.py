from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = ROOT / "deliverables" / "GEO写作检查清单.docx"

TITLE = "GEO 写作检查清单"
SUBTITLE = "版本：v1.0    整理日期：2026-04-22"

INTRO = [
    "本清单适用于品牌站博客、专题页、FAQ 页、产品支持型内容等 GEO 写作场景。",
    "GEO 的核心目标不是单纯堆关键词，而是让内容更容易被 AI 系统理解、抽取、复述、引用和推荐。",
    "建议在选题和提纲、初稿完成后、发布前三个阶段分别使用本清单。",
]

SECTIONS = [
    (
        "一、写前准备",
        [
            "已明确文章核心问题，且可用一句话概括。",
            "已确认文章主要搜索意图：信息型 / 比较型 / 决策型 / 转化型。",
            "已确认目标国家或站点，不混用不同市场规则。",
            "已确认核心产品、配件、功能、政策或场景实体。",
            "已确认品牌官方页面、产品页、支持页、帮助中心页等可信来源。",
            "已确认文中涉及的产品名、型号名、容量、功率、兼容性、保修信息。",
            "已提前确认哪些信息可以写，哪些信息不能写或不能绝对化表达。",
        ],
    ),
    (
        "二、内容结构检查",
        [
            "标题直接对应用户问题，不空泛，不偏题。",
            "开头 100-150 词内直接回答问题，不长篇铺垫。",
            "首段已明确：结论是什么、适合谁、前提条件是什么。",
            "正文结构清晰，至少包含：结论、解释、比较或场景、FAQ。",
            "各级标题围绕用户追问展开，而不是围绕作者表达习惯展开。",
            "文章中段有可快速扫描的结构化模块，如表格、对比块、步骤块、条件列表。",
            "结尾有总结，不重复堆砌卖点。",
        ],
    ),
    (
        "三、可引用性检查",
        [
            "文中存在可被 AI 直接摘取的结论句。",
            "关键结论使用完整句表达，不依赖上下文才能理解。",
            "同一问题下的核心答案集中，不分散在多个段落中。",
            "重要参数、结论或差异点优先使用列表或表格呈现。",
            "比较类内容明确给出比较维度，而非只写主观评价。",
            "对适合谁、不适合谁、适用场景、限制条件有明确表达。",
        ],
    ),
    (
        "四、事实与证据检查",
        [
            "所有关键数据均可追溯到官方页面、帮助中心、说明书或可信机构。",
            "没有将猜测、经验判断或推断写成确定事实。",
            "没有擅自补充官网未明确披露的认证、测试或兼容性结论。",
            "涉及时间敏感内容时，已确认当前国家、政策、价格或规则仍有效。",
            "涉及参数时，已区分清楚容量、输出功率、输入功率、可扩容范围等概念。",
            "若存在限制条件，已在正文显式写出，而不是只写优点。",
            "若无法验证的内容必须保留，已使用谨慎表述并标注前提。",
        ],
    ),
    (
        "五、品牌、产品与合规检查",
        [
            "产品名、系列名、型号名全文保持一致。",
            "品牌名拼写统一，国家、地区、政策名写法统一。",
            "未使用客户禁词、违规承诺、夸大表述或绝对化说法。",
            "没有把产品能力写成超出官方页面范围的结论。",
            "涉及兼容性、安装、补贴、保修、认证、节省金额等内容时，已加入必要限制说明。",
            "涉及法规、税务、安装建议时，未替代专业意见。",
            "如客户要求免责声明、前置说明或特定段落，已按位置要求添加。",
        ],
    ),
    (
        "六、内链与页面信号检查",
        [
            "首屏或前两段已自然加入重点内链。",
            "内链优先链接到官方产品页、帮助页、对比页或专题页。",
            "没有使用过时、失效或年份错误的链接。",
            "锚文本具备语义，不只使用点击这里或纯网址。",
            "推荐产品与正文主题强相关，不生硬插入。",
            "页面中包含 FAQ、表格、摘要块等易被 AI 理解的结构。",
            "如页面支持，已考虑 FAQ Schema、Article Schema、Product Schema、作者信息、更新时间等信号。",
        ],
    ),
    (
        "七、FAQ 检查",
        [
            "FAQ 问题来自真实追问，而不是改写正文小标题。",
            "FAQ 回答简洁直接，先回答再补充说明。",
            "FAQ 与正文不机械重复，能补足用户后续问题。",
            "FAQ 涵盖适用性、差异、限制、兼容性、安装、价格、保修等高频问题。",
            "FAQ 中的答案仍保持事实可验证，不引入新风险。",
        ],
    ),
    (
        "八、发布前质检",
        [
            "已重新检查标题、H1、摘要和首段是否一致。",
            "已检查文章是否真正回答了目标关键词，而不是只提到了关键词。",
            "已检查是否存在废话开头、空泛过渡和营销腔堆砌。",
            "已检查是否存在参数错误、实体混乱、型号混写或语言不统一。",
            "已检查所有链接是否可打开且落地页正确。",
            "已检查图片说明、图表标题、模块名称是否与正文一致。",
            "已检查全文是否便于复制、摘录、总结和引用。",
        ],
    ),
]

VETO_ITEMS = [
    "核心结论无法在前文直接找到。",
    "参数、兼容性、保修、政策等关键信息无法验证。",
    "标题和正文回答的问题不一致。",
    "使用了客户明确禁止的说法。",
    "必要免责声明缺失。",
    "内链错误或推荐产品与主题明显不符。",
    "FAQ 重复正文且没有新增价值。",
]

TIPS = [
    "先给答案，再给原因。",
    "多写适合谁、不适合谁、在什么条件下成立。",
    "多用表格、对比块、步骤块、条件块。",
    "结论句尽量完整，便于 AI 直接复述。",
    "关键数据后最好自然体现来源依据，如 based on official specifications。",
    "如果信息只在特定国家、特定页面、特定条件下成立，务必写清前提。",
]

REVIEW_ORDER = [
    "先看标题和首段是否直接回答问题。",
    "再看实体、参数、兼容性和免责声明是否正确。",
    "再看中段是否有结构化信息可供引用。",
    "再看 FAQ 是否承接追问。",
    "最后检查内链、页面信号和整体可读性。",
]


def set_document_defaults(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")

    for style_name in ("Title", "Heading 1", "Heading 2"):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")


def add_checkbox_items(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        run = paragraph.add_run(f"□ {item}")
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")


def build_document() -> Document:
    document = Document()
    set_document_defaults(document)

    section = document.sections[0]
    section.top_margin = Pt(48)
    section.bottom_margin = Pt(48)
    section.left_margin = Pt(54)
    section.right_margin = Pt(54)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(SUBTITLE)
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(10.5)
    subtitle_run.font.name = "Calibri"
    subtitle_run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")

    document.add_paragraph("")

    document.add_heading("使用说明", level=1)
    for line in INTRO:
        document.add_paragraph(line)

    for heading, items in SECTIONS:
        document.add_heading(heading, level=1)
        add_checkbox_items(document, items)

    document.add_heading("九、一票否决项", level=1)
    document.add_paragraph("以下问题如未解决，不建议发布：")
    add_checkbox_items(document, VETO_ITEMS)

    document.add_heading("十、推荐写法提示", level=1)
    add_checkbox_items(document, TIPS)

    document.add_heading("十一、推荐审核顺序", level=1)
    for index, item in enumerate(REVIEW_ORDER, start=1):
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"{index}. {item}")
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")

    document.add_section(WD_SECTION.NEW_PAGE)
    last_paragraph = document.add_paragraph("备注：本清单可作为 GEO 写手自检表，也可作为编辑、审核或客户复核口径。")
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    return document


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
